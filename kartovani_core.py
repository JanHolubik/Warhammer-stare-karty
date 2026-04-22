from __future__ import annotations

from pathlib import Path
from io import BytesIO
from typing import Dict
import re
import glob

import pandas as pd
from docx import Document
from bs4 import BeautifulSoup

VAT_RATE = 1.21


def load_color_catalog() -> dict[str, str]:
    """
    Načte products-*.csv ze složky colours a vrátí mapu:
    normalized_color_name -> code_or_ean
    """
    base_dir = Path(__file__).resolve().parent
    csv_files = sorted(glob.glob(str(base_dir / "colours/products-*.csv")))

    color_map: dict[str, str] = {}

    for csv_file in csv_files:
        try:
            df = pd.read_csv(csv_file, sep=";", dtype=str).fillna("")

            for _, row in df.iterrows():
                product_name = str(row.get("name", "")).strip()
                product_code = str(row.get("code", "")).strip()
                product_ean = str(row.get("ean", "")).strip()

                if not product_name:
                    continue

                target_value = product_code or product_ean
                if not target_value:
                    continue

                normalized = normalize_color_name(product_name)
                if normalized:
                    color_map[normalized] = target_value

        except Exception:
            continue

    return color_map


def normalize_color_name(text: str) -> str:
    text = str(text or "").strip().lower()

    prefixes = [
        "citadel base:",
        "citadel layer:",
        "citadel shade:",
        "citadel contrast:",
        "citadel dry:",
        "citadel technical:",
        "citadel spray:",
        "citadel air:",
        "citadel colour:",
    ]

    for prefix in prefixes:
        if text.startswith(prefix):
            text = text[len(prefix):].strip()

    text = re.sub(r"\s+", " ", text)
    return text


def extract_related_products_from_text(
    text: str,
    color_map: dict[str, str],
    limit: int = 7,
) -> list[str]:
    """
    Najde barvy v textu podle názvu a vrátí jejich kódy v pořadí výskytu.
    """
    source_text = str(text or "")
    source_text_lower = source_text.lower()

    matches: list[tuple[int, str]] = []

    for color_name, product_code in color_map.items():
        if not color_name:
            continue

        pattern = r"(?<!\w)" + re.escape(color_name.lower()) + r"(?!\w)"
        match = re.search(pattern, source_text_lower, flags=re.IGNORECASE)

        if match:
            matches.append((match.start(), product_code))

    matches.sort(key=lambda x: x[0])

    seen = set()
    result: list[str] = []

    for _, product_code in matches:
        if product_code in seen:
            continue
        seen.add(product_code)
        result.append(product_code)

        if len(result) >= limit:
            break

    return result


def enrich_kartovani_df_with_related_products(
    df: pd.DataFrame,
    row_index: int,
    limit: int = 13,
) -> pd.DataFrame:
    """
    Z výsledného HTML popisu vytáhne názvy barev a doplní je do relatedProduct sloupců.
    """
    color_map = load_color_catalog()
    if not color_map:
        return df

    df_out = df.copy()

    related_columns = ["relatedProduct"] + [f"relatedProduct{i}" for i in range(2, limit + 1)]
    for col in related_columns:
        if col not in df_out.columns:
            df_out[col] = ""

    combined_text = "\n".join(
        [
            str(df_out.at[row_index, "description:cs"]) if "description:cs" in df_out.columns else "",
            str(df_out.at[row_index, "description:en"]) if "description:en" in df_out.columns else "",
            str(df_out.at[row_index, "description:sk"]) if "description:sk" in df_out.columns else "",
            str(df_out.at[row_index, "description"]) if "description" in df_out.columns else "",
        ]
    )

    combined_text = BeautifulSoup(combined_text, "html.parser").get_text(" ", strip=True)

    related_codes = extract_related_products_from_text(
        text=combined_text,
        color_map=color_map,
        limit=limit,
    )

    for i, col in enumerate(related_columns):
        df_out.at[row_index, col] = related_codes[i] if i < len(related_codes) else ""

    return df_out


BASE_DIR = Path(__file__).resolve().parent
PROMPT_TEMPLATE_DIR = BASE_DIR / "prompt_templates"
TEMPLATE_DIR = BASE_DIR / "sablony"

INVALID_IMAGE_VALUES = {
    "", "/", "placeholder",
    "intro_image_src", "/intro_image_src",
    "img1_src", "/img1_src",
    "img2_src", "/img2_src",
    "img3_src", "/img3_src",
    "img4_src", "/img4_src",
    "img5_src", "/img5_src",
    "img6_src", "/img6_src",
    "img7_src", "/img7_src",
    "img8_src", "/img8_src",
    "img9_src", "/img9_src",
    "img10_src", "/img10_src",
}

INVALID_VIDEO_VALUES = {
    "", "/", "https://www.youtube.com/embed/", "https://www.youtube.com/embed"
}


def make_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def pick_existing(template_dir: Path, candidates: list[str]) -> Path:
    for name in candidates:
        path = template_dir / name
        if path.exists():
            return path
    raise FileNotFoundError(
        f"Nenalezena žádná z těchto šablon ve složce {template_dir}: {candidates}"
    )


def get_template_paths(template_dir: Path, template_type: str) -> Dict[str, Path]:
    template_type = template_type.strip().lower()

    short_map = {
        "miniatures": ["kratky.docx"],
        "books": ["kratky_univ.docx"],
        "warscroll": ["kratky_univ.docx"],
        "dice": ["kratky.docx"],
        "upgrades": ["kratky.docx"],
        "accessories": ["stetce_kratky_text.docx"],
        "albi": ["kratky_albi.docx"],
    }

    detail_map = {
        "miniatures": ["detailni.docx"],
        "books": ["detailni_kniha.docx"],
        "warscroll": ["detailni_warscoll.docx", "detailni_warscroll.docx"],
        "dice": ["detailni_kostky.docx"],
        "upgrades": ["detailni_upgrades.docx"],
        "accessories": ["stetce_pris.docx"],
        "albi": ["detailni_albi.docx"],
    }

    short_map_en = {
        "miniatures": ["kratky_en.docx"],
        "books": ["kratky_univ_en.docx"],
        "warscroll": ["kratky_univ_en.docx"],
        "dice": ["kratky_en.docx"],
        "upgrades": ["kratky_en.docx"],
        "accessories": ["stetce_kratky_text_en.docx"],
        "albi": ["kratky_albi_en.docx"],
    }

    detail_map_en = {
        "miniatures": ["detailni_en.docx"],
        "books": ["detailni_kniha_en.docx"],
        "warscroll": ["detailni_warscoll_en.docx", "detailni_warscroll_en.docx"],
        "dice": ["detailni_kostky_en.docx"],
        "upgrades": ["detailni_upgrades_en.docx"],
        "accessories": ["stetce_pris en.docx"],
        "albi": ["detailni_albi_en.docx"],
    }

    short_map_sk = {
        "miniatures": ["kratky_sk.docx"],
        "books": ["kratky_univ_sk.docx"],
        "warscroll": ["kratky_univ_sk.docx"],
        "dice": ["kratky_sk.docx"],
        "upgrades": ["kratky_sk.docx"],
        "accessories": ["stetce_kratky_text_sk.docx"],
        "albi": ["kratky_albi_sk.docx"],
    }

    detail_map_sk = {
        "miniatures": ["detailni_sk.docx"],
        "books": ["detailni_kniha_sk.docx"],
        "warscroll": ["detailni_warscoll_sk.docx", "detailni_warscroll_sk.docx"],
        "dice": ["detailni_kostky_sk.docx"],
        "upgrades": ["detailni_upgrades_sk.docx"],
        "accessories": ["stetce_pris sk.docx"],
        "albi": ["detailni_albi_sk.docx"],
    }

    if template_type not in short_map:
        raise ValueError(f"Neznámý TEMPLATE_TYPE: {template_type}")

    return {
        "short_cs": pick_existing(template_dir, short_map[template_type]),
        "detail_cs": pick_existing(template_dir, detail_map[template_type]),
        "short_en": pick_existing(template_dir, short_map_en[template_type]),
        "detail_en": pick_existing(template_dir, detail_map_en[template_type]),
        "short_sk": pick_existing(template_dir, short_map_sk[template_type]),
        "detail_sk": pick_existing(template_dir, detail_map_sk[template_type]),
    }


def create_kartovani_card_row(
    name: str,
    code: str,
    external_code: str,
    ean: str,
    price: float,
    product_type: str,
    description: str = "",
    image_urls: list[str] | None = None,
) -> pd.DataFrame:
    image_urls = image_urls or []
    price_without_vat = round(float(price) / VAT_RATE, 2) if price else 0.0

    row = {
        "code": code,
        "pairCode": "",
        "externalCode": external_code.strip() or code,
        "name": name,
        "ean": ean,
        "price": price,
        "priceWithoutVat": price_without_vat,
        "description": description,
        "manufacturer": "",
        "availabilityInStock": "Skladem",
        "availabilityOutOfStock": "Na dotaz",
        "itemType": product_type,
        "productVisibility": "visible",
        "relatedVideo": "",
    }

    for idx in range(1, 11):
        col = "image" if idx == 1 else f"image{idx}"
        row[col] = image_urls[idx - 1] if idx <= len(image_urls) else ""

    return pd.DataFrame([row])


def create_kartovani_source_row(
    name: str,
    code: str,
    external_code: str,
    ean: str,
    price: float,
    intro_image_src: str = "",
    image_urls: list[str] | None = None,
    video_url: str = "",
    template_kind: str = "",
    prompt_type: str = "",
) -> pd.DataFrame:
    image_urls = image_urls or []
    price_without_vat = round(float(price) / VAT_RATE, 2) if price else 0.0

    if not template_kind:
        raise ValueError("template_kind musí být vyplněný (např. 'miniatures', 'accessories', 'books')")
    if not prompt_type:
        raise ValueError("prompt_type musí být vyplněný (např. 'miniatures', 'accessories', 'books')")

    row = {
        "code": code,
        "pairCode": "",
        "externalCode": external_code.strip() or code,
        "product_name": name,
        "name": name,
        "ean": ean,
        "price": price,
        "priceWithoutVat": price_without_vat,
        "description": "",
        "manufacturer": "",
        "availabilityInStock": "Skladem",
        "availabilityOutOfStock": "Na dotaz",
        "itemType": "product",
        "intro_image_src": intro_image_src,
        "productVisibility": "visible",
        "video_url": video_url,
        "template_kind": template_kind,
        "prompt_type": prompt_type,
    }

    for idx in range(1, 11):
        row[f"img{idx}_src"] = image_urls[idx - 1] if idx <= len(image_urls) else ""
        row[f"image_alt_{idx}"] = ""

        target_col = "image" if idx == 1 else f"image{idx}"
        row[target_col] = image_urls[idx - 1] if idx <= len(image_urls) else ""

    for lang in ["cs", "en", "sk"]:
        row[f"name:{lang}"] = ""
        row[f"shortDescription:{lang}"] = ""
        row[f"description:{lang}"] = ""
        row[f"seoTitle:{lang}"] = ""
        row[f"metaDescription:{lang}"] = ""
        row[f"xmlFeedName:{lang}"] = ""

    return pd.DataFrame([row])


def build_kartovani_prompt(
    prompt_type: str,
    product_name: str,
    product_ean: str,
    product_code: str = "",
    lang: str = "cs",
) -> str:
    lang = lang.strip().lower()

    if lang not in {"cs", "en", "sk"}:
        raise ValueError(f"Neplatný jazyk: {lang}")

    template_path = PROMPT_TEMPLATE_DIR / f"{prompt_type}_{lang}.txt"
    if not template_path.exists():
        raise FileNotFoundError(f"Šablona promptu nenalezena: {template_path}")

    template_text = template_path.read_text(encoding="utf-8")

    return f"""{template_text}

--------------------------------------------------
SEQUENTIAL API MODE
--------------------------------------------------

V tomto běhu generuj pouze 1 jazyk.

LANG_MODE
{lang}

PRAVIDLO
Vrať pouze jeden jazykový blok:
[LANG={lang}]

Nikdy negeneruj jiné jazykové bloky než [LANG={lang}].

--------------------------------------------------
PRODUKT
{product_name}

EAN
{product_ean}

CODE
{product_code}
--------------------------------------------------
"""


def parse_ai_output_to_lang_blocks(text: str) -> dict:
    pattern = r"\[LANG=(cs|en|sk)\]\s*(.*?)(?=\[LANG=cs\]|\[LANG=en\]|\[LANG=sk\]|\Z)"
    matches = re.findall(pattern, text, flags=re.DOTALL | re.IGNORECASE)

    out = {"cs": "", "en": "", "sk": ""}
    for lang, content in matches:
        out[lang.lower()] = content.strip()
    return out


def parse_key_value_block(block_text: str) -> dict:
    result = {}
    current_key = None
    current_value_lines = []

    for line in block_text.splitlines():
        if ":" in line:
            maybe_key, maybe_value = line.split(":", 1)
            key = maybe_key.strip()
            if key:
                if current_key is not None:
                    result[current_key] = "\n".join(current_value_lines).strip()
                current_key = key
                current_value_lines = [maybe_value.strip()]
                continue

        if current_key is not None:
            current_value_lines.append(line.strip())

    if current_key is not None:
        result[current_key] = "\n".join(current_value_lines).strip()

    return result


def replace_placeholders_in_docx(template_path: Path, values: dict) -> str:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, value in values.items():
            new_text = new_text.replace("{" + key + "}", value or "")
        if new_text != full_text:
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            paragraph.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    new_text = full_text
                    for key, value in values.items():
                        new_text = new_text.replace("{" + key + "}", value or "")
                    if new_text != full_text:
                        for i in range(len(paragraph.runs) - 1, -1, -1):
                            paragraph._element.remove(paragraph.runs[i]._element)
                        paragraph.add_run(new_text)

    html_lines = []

    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt:
            html_lines.append(txt)

    for table in doc.tables:
        html_lines.append("<table>")
        for row in table.rows:
            html_lines.append("<tr>")
            for cell in row.cells:
                html_lines.append(f"<td>{cell.text.strip()}</td>")
            html_lines.append("</tr>")
        html_lines.append("</table>")

    return "\n".join(html_lines)


def normalize_miniature_values(values: dict) -> dict:
    normalized = dict(values)

    if "nazev_produktu" in normalized:
        normalized.setdefault("název_sady", normalized["nazev_produktu"])
        normalized.setdefault("nazev_sady", normalized["nazev_produktu"])

    if "pro_koho" in normalized:
        normalized.setdefault("doporučení_pro_koho", normalized["pro_koho"])
        normalized.setdefault("doporuceni_pro_koho", normalized["pro_koho"])

    if "potrebne_vybaveni" in normalized:
        normalized.setdefault("potřebné_vybavení", normalized["potrebne_vybaveni"])

    if "doporuceni_pro_koho" in normalized:
        normalized.setdefault("doporučení_pro_koho", normalized["doporuceni_pro_koho"])

    return normalized


def is_valid_image(value: str | None) -> bool:
    if value is None:
        return False
    return str(value).strip() not in INVALID_IMAGE_VALUES


def is_valid_video(value: str | None) -> bool:
    if value is None:
        return False
    return str(value).strip() not in INVALID_VIDEO_VALUES


def add_class(tag, class_name: str) -> None:
    classes = tag.get("class", [])
    if class_name not in classes:
        classes.append(class_name)
        tag["class"] = classes


def make_images_clickable(soup: BeautifulSoup) -> None:
    for img in soup.find_all("img"):
        src = (img.get("src") or "").strip()
        if not is_valid_image(src):
            continue

        parent = img.parent
        if parent and parent.name == "a" and "image-link" in (parent.get("class") or []):
            continue

        a = soup.new_tag(
            "a",
            href=src,
            target="_blank",
            rel="noopener noreferrer"
        )
        a["class"] = ["image-link"]
        img.wrap(a)


def cleanup_rendered_html(html: str, values: dict) -> str:
    soup = BeautifulSoup(html, "html.parser")

    intro_valid = is_valid_image(values.get("intro_image_src", ""))
    img1_valid = is_valid_image(values.get("img1_src", ""))
    img2_valid = is_valid_image(values.get("img2_src", ""))
    img3_valid = is_valid_image(values.get("img3_src", ""))
    img4_valid = is_valid_image(values.get("img4_src", ""))
    video_valid = is_valid_video(values.get("video_url", ""))

    hero_image = soup.select_one(".hero-image")
    hero = soup.select_one(".hero")
    if hero_image and not intro_valid:
        hero_image.decompose()
        if hero:
            add_class(hero, "no-image")

    story_media = soup.select_one(".story-media")
    story_split = soup.select_one(".story-split")
    if story_media and not img1_valid:
        story_media.decompose()
        if story_split:
            add_class(story_split, "no-image")

    full_bleed = soup.select_one(".full-bleed-image")
    if full_bleed and not img2_valid:
        prev = full_bleed.find_previous_sibling()
        if prev and "soft-divider" in (prev.get("class") or []):
            prev.decompose()
        full_bleed.decompose()

    duo_gallery = soup.select_one(".duo-gallery")
    if duo_gallery:
        duo_cards = duo_gallery.select(":scope > .duo-card")
        if len(duo_cards) >= 1 and not img3_valid:
            duo_cards[0].decompose()

        duo_cards = duo_gallery.select(":scope > .duo-card")
        if len(duo_cards) >= 2 and not img4_valid:
            duo_cards[1].decompose()

        duo_cards = duo_gallery.select(":scope > .duo-card")
        if len(duo_cards) == 1:
            add_class(duo_gallery, "one-item")
        elif len(duo_cards) == 0:
            duo_gallery.decompose()

    if not video_valid:
        for section in soup.select("section.section-card"):
            iframe = section.select_one("iframe")
            if iframe:
                section.decompose()

    for img in soup.find_all("img"):
        src = (img.get("src") or "").strip()
        if not is_valid_image(src):
            parent = img.parent
            if parent and parent.name == "a" and "image-link" in (parent.get("class") or []):
                parent.decompose()
            else:
                img.decompose()

    for iframe in soup.find_all("iframe"):
        src = (iframe.get("src") or "").strip()
        if not is_valid_video(src):
            section = iframe.find_parent("section")
            if section:
                section.decompose()
            else:
                iframe.decompose()

    make_images_clickable(soup)

    return str(soup)


def validate_final_html(html: str) -> list[str]:
    forbidden = [
        'src="intro_image_src"',
        'src="/intro_image_src"',
        'src="img1_src"',
        'src="/img1_src"',
        'src="img2_src"',
        'src="/img2_src"',
        'src="img3_src"',
        'src="/img3_src"',
        'src="img4_src"',
        'src="/img4_src"',
        'src="img5_src"',
        'src="/img5_src"',
        'src="img6_src"',
        'src="/img6_src"',
        'src="img7_src"',
        'src="/img7_src"',
        'src="img8_src"',
        'src="/img8_src"',
        'src="img9_src"',
        'src="/img9_src"',
        'src="img10_src"',
        'src="/img10_src"',
        'src=""',
        'src="/"',
        'https://www.youtube.com/embed/"',
    ]
    return [x for x in forbidden if x in html]


def build_kartovani_html(ai_output: str, template_kind: str, extra_values: dict | None = None) -> dict:
    lang_blocks = parse_ai_output_to_lang_blocks(ai_output)

    template_paths = get_template_paths(TEMPLATE_DIR, template_kind)

    short_files = {
        "cs": template_paths["short_cs"],
        "en": template_paths["short_en"],
        "sk": template_paths["short_sk"],
    }
    long_files = {
        "cs": template_paths["detail_cs"],
        "en": template_paths["detail_en"],
        "sk": template_paths["detail_sk"],
    }

    out = {}

    for lang in ["cs", "en", "sk"]:
        values = parse_key_value_block(lang_blocks.get(lang, ""))

        if extra_values:
            values.update(extra_values)

        if not values.get("intro_image_src", "").strip():
            if values.get("img1_src", "").strip():
                values["intro_image_src"] = values["img1_src"]

        if template_kind == "miniatures":
            values = normalize_miniature_values(values)

        short_html = replace_placeholders_in_docx(short_files[lang], values)
        long_html = replace_placeholders_in_docx(long_files[lang], values)

        long_html = cleanup_rendered_html(long_html, values)

        short_errors = validate_final_html(short_html)
        long_errors = validate_final_html(long_html)

        if short_errors:
            raise ValueError(f"Neplatné placeholdery v shortDescription:{lang}: {short_errors}")
        if long_errors:
            raise ValueError(f"Neplatné placeholdery v description:{lang}: {long_errors}")

        out[f"shortDescription:{lang}"] = short_html
        out[f"description:{lang}"] = long_html

        product_name = (
            values.get("nazev_produktu", "").strip()
            or values.get("název_sady", "").strip()
            or values.get("nazev_sady", "").strip()
        )
        short_desc = (
            values.get("kratky_popis", "").strip()
            or values.get("strucny_popis_produktu", "").strip()
        )

        final_short = f"{product_name} {short_desc}".strip()
        final_short = final_short.replace("\n", " ").replace("\r", " ")

        out[f"name:{lang}"] = product_name
        out[f"seoTitle:{lang}"] = product_name
        out[f"xmlFeedName:{lang}"] = product_name
        out[f"metaDescription:{lang}"] = final_short[:155] if final_short else ""

    return out


def apply_kartovani_output_to_csv(
    df: pd.DataFrame,
    row_index: int,
    ai_output: str,
    template_kind: str,
    extra_values: dict | None = None,
) -> pd.DataFrame:
    html_map = build_kartovani_html(ai_output, template_kind, extra_values=extra_values)
    df_out = df.copy()

    for col, value in html_map.items():
        if col not in df_out.columns:
            df_out[col] = ""
        df_out.at[row_index, col] = value

    if extra_values and "video_url" in extra_values:
        if "video_url" not in df_out.columns:
            df_out["video_url"] = ""
        df_out.at[row_index, "video_url"] = extra_values["video_url"]

    df_out = df_out.drop(
        columns=[c for c in ["product_name", "name", "description", "shortDescription"] if c in df_out.columns],
        errors="ignore"
    )

    preferred_order = [
        "code",
        "pairCode",
        "externalCode",
        "name:cs",
        "name:en",
        "name:sk",
        "shortDescription:cs",
        "shortDescription:en",
        "shortDescription:sk",
        "description:cs",
        "description:en",
        "description:sk",
        "price",
        "priceWithoutVat",
        "manufacturer",
        "itemType",
        "video_url",
        "relatedVideo",
        "availabilityInStock",
        "availabilityOutOfStock",
        "ean",
        "productVisibility",
        "xmlFeedName:cs",
        "xmlFeedName:en",
        "xmlFeedName:sk",
        "seoTitle:cs",
        "seoTitle:en",
        "seoTitle:sk",
        "metaDescription:cs",
        "metaDescription:en",
        "metaDescription:sk",
        "template_kind",
        "prompt_type",
    ]

    existing_preferred = [col for col in preferred_order if col in df_out.columns]
    remaining_cols = [col for col in df_out.columns if col not in existing_preferred]

    df_out = df_out[existing_preferred + remaining_cols]
    df_out = df_out.fillna("")

    df_out = enrich_kartovani_df_with_related_products(
        df=df_out,
        row_index=row_index,
        limit=13,
    )

    return df_out